import os
from django.shortcuts import render
from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.http import HttpResponseRedirect
from .models import *
from .models import ManualInput
from docx import Document
import sqlite3 as sq
from django.shortcuts import render
import matplotlib.pyplot as plt
from docx.shared import Inches
import glob
from django.contrib.auth.decorators import permission_required

def CreateGraphTcelllink(neu_cd3, neu_cd4, neu_cd8, neu_lymf, surname, date):

    x_max = [0.0, -1.67, 0.0, 3.00, 0.0]
    x_min = [0.0, -1.80, 0.0, 5.00, 0.0]
    x_pac = [0.0, -neu_lymf, 0.0, neu_cd8, 0.0]

    y_max = [2.30, 0.0, -9.50, 0.0, 2.30]
    y_min = [3.70, 0.0, -12.30, 0.0, 3.70]
    y_pac = [neu_cd3, 0.0, -neu_cd4, 0.0, neu_cd3]

    plt.plot(x_max, y_max, linestyle='--', color='green')
    plt.plot(x_min, y_min, linestyle='--', color='green')
    plt.plot(x_pac, y_pac, linestyle='-', color='red')

    plt.scatter(x_max, y_max, color='green')
    plt.scatter(x_min, y_min, color='green')
    plt.scatter(x_pac, y_pac, color='red')

    plt.annotate('NEU/СD3', xy=(0.0, neu_cd3), xytext=(0.0, neu_cd3))
    plt.annotate('NEU/LYMF', xy=(-neu_lymf, 0.0), xytext=(-neu_lymf, 0.0))
    plt.annotate('NEU/CD4', xy=(0.0, neu_cd4), xytext=(0.0, -neu_cd4))
    plt.annotate('NEU/CD8', xy=(neu_cd8, 0.0), xytext=(neu_cd8, 0.0))
    plt.grid(True)
    # добавление заголовка и меток осей
    plt.title('График Т-клеточного звена')
    plt.xlabel('Ось x')
    plt.ylabel('Ось y')

    # Отображение графика
    plt.savefig(f'Графики/T-клеточный {surname} {date}.png')
    plt.close()

def CreateGraphBcelllink(neu_lymf, neu_cd19, cd19_cd4, cd19_cd8, surname, date):

    x_max = [0, -1.67, 0, 0.20, 0]
    x_min = [0, -1.80, 0, 0.30, 0]
    x_pac = [0, -neu_lymf, 0, cd19_cd8, 0]

    y_max = [9.60, 0, -0.50, 0, 9.60]
    y_min = [10.00, 0, -0.80, 0, 10.00]
    y_pac = [neu_cd19, 0, -cd19_cd4, 0, neu_cd19]

    plt.plot(x_max, y_max, linestyle='--', color='green')
    plt.plot(x_min, y_min, linestyle='--', color='green')
    plt.plot(x_pac, y_pac, linestyle='-', color='red')

    plt.scatter(x_max, y_max, color='green')
    plt.scatter(x_min, y_min, color='green')
    plt.scatter(x_pac, y_pac, color='red')

    plt.annotate('NEU/LYMF', xy=(0.0, neu_cd19), xytext=(0.0, neu_cd19))
    plt.annotate('NEU/CD19', xy=(-neu_lymf, 0.0), xytext=(-neu_lymf, 0.0))
    plt.annotate('CD19/CD4', xy=(0.0, -cd19_cd4), xytext=(0.0, -cd19_cd4))
    plt.annotate('CD19/CD8', xy=(cd19_cd8, 0.0), xytext=(cd19_cd8, 0.0))
    plt.grid(True)





    # добавление заголовка и меток осей
    plt.title('График В-клеточного звена')
    plt.xlabel('Ось x')
    plt.ylabel('Ось y')

    # Отображение графика
    plt.savefig(f'Графики/B-клеточный {surname} {date}.png')
    plt.close()

def CreateGraphCytokinePairs(interlikin, interferon, fno, surname, date):
    fig = plt.figure()
    ax = fig.add_subplot(111, projection='3d')

    x = [120, 0, 0, 120]
    y = [0, 120, 0, 0]
    z = [0, 0, 120, 0]

    x1 = [80, 0, 0, 80]
    y2 = [0, 80, 0, 0]
    z3 = [0, 0, 80, 0]

    x11 = [interferon, 0, 0, interferon]
    y21 = [0,interlikin, 0, 0]
    z31 = [0, 0, fno, 0]

    ax.plot(x, y, z, linestyle='--', color='green')
    ax.plot(x1, y2, z3, linestyle='--', color='green')
    ax.plot(x11, y21, z31, linestyle='-', color='red')

    ax.set_xlabel('X')
    ax.set_ylabel('Y')
    ax.set_zlabel('Z')
    ax.view_init(elev=50, azim=43)
    #ax.set_axis_off()
    ax.text(x=140, y=0, z=0, s="Интерферон", fontsize=10)
    ax.text(x=0, y=145, z=0, s="Интерликин", fontsize=10)
    ax.text(x=0, y=0, z=125, s="ФНО", fontsize=10)
    #ax.tick_params(axis='both', which='both', labelcolor='none')
    plt.savefig(f'Графики/Цитокиновые пары {surname} {date}.png')
    plt.close()


def index(request):
    user = User.objects.filter(id=request.user.id)
    if len(user) != 0:
        return render(request, 'index.html')
    else:
        return redirect('login')

def calculate(request):
    user = User.objects.filter(id=request.user.id)
    if len(user) != 0:
        return render(request, 'calculate.html')
    else:
        return redirect('login')

def docxForms(request):
    user = User.objects.filter(id=request.user.id)
    file_path = 'Печатные формы\*.docx'
    files = glob.glob(file_path)

    context = {
        'files': files
    }
    if len(user) != 0:
        return render(request, 'docxForms.html', context)
    else:
        return redirect('login')

# Import mimetypes module
import mimetypes
# import os module
import os
# Import HttpResponse module
from django.http.response import HttpResponse

def test(request):
    user = User.objects.filter(id=request.user.id)
    if len(user) != 0:
        return render(request, 'test.html')
    else:
        return redirect('login')

def create(request):
    if request.method == "POST":
        man_in = ManualInput()
        man_in.name = request.POST.get('name')
        man_in.manwom = request.POST.get('manwom')
        man_in.analysis_date = request.POST.get('analysis_date')
        man_in.age = request.POST.get('age')
        man_in.main_diagnosis = request.POST.get('main_diagnosis')
        man_in.co_diagnosis = request.POST.get('co_diagnosis')
        man_in.genes1 = request.POST.get('genes1')
        man_in.genes2 = request.POST.get('genes2')
        man_in.genes3 = request.POST.get('genes3')
        man_in.season = request.POST.get('season')
        man_in.leu_wbc = request.POST.get('leu_wbc')
        man_in.lymf = request.POST.get('lymf')
        man_in.mon = request.POST.get('mon')
        man_in.neu = request.POST.get('neu')
        man_in.hgb = request.POST.get('hgb')
        man_in.trom_plt = request.POST.get('trom_plt')
        man_in.eos = request.POST.get('eos')
        man_in.bas = request.POST.get('bas')
        man_in.b_lymf = request.POST.get('b_lymf')
        man_in.t_lymf = request.POST.get('t_lymf')
        man_in.t_help = request.POST.get('t_help')
        man_in.t_cyto_lymf = request.POST.get('t_cyto_lymf')
        man_in.cd3_cd4_cd3_cd8 = request.POST.get('cd3_cd4_cd3_cd8')
        man_in.nk_cells_cyto = request.POST.get('nk_cells_cyto')
        man_in.common_nk_cells = request.POST.get('common_nk_cells')
        man_in.circ_immu_comps = request.POST.get('circ_immu_comps')
        man_in.hct_test_spon = request.POST.get('hct_test_spon')
        man_in.hct_test_stim = request.POST.get('hct_test_stim')
        man_in.cd3_ifny_stim = request.POST.get('cd3_ifny_stim')
        man_in.cd3_ifny_spon = request.POST.get('cd3_ifny_spon')
        man_in.cd3_tnfa_stim = request.POST.get('cd3_tnfa_stim')
        man_in.cd3_tnfa_spon = request.POST.get('cd3_tnfa_spon')
        man_in.cd3_il2_stim = request.POST.get('cd3_il2_stim')
        man_in.cd3_il2_spon = request.POST.get('cd3_il2_spon')
        man_in.fno = request.POST.get('fno')
        man_in.save()
    return HttpResponseRedirect("/")
@permission_required('analysis.view_recomendation', raise_exception=True)
@permission_required('analysis.add_recomendation', raise_exception=True)
def createRec(request):
    if request.method == "POST":
        rec = Recomendation()
        rec.age = request.POST.get('age')
        rec.season = request.POST.get('season')

        rec.leu_wbc = request.POST.get('leu_wbc')
        rec.leu_wbc_min = request.POST.get('leu_wbc_min')
        rec.leu_wbc_max = request.POST.get('leu_wbc_max')

        rec.lymf = request.POST.get('lymf')
        rec.lymf_min = request.POST.get('lymf_min')
        rec.lymf_max = request.POST.get('lymf_max')

        rec.mon = request.POST.get('mon')
        rec.mon_min = request.POST.get('mon_min')
        rec.mon_max = request.POST.get('mon_max')

        rec.neu = request.POST.get('neu')
        rec.neu_min = request.POST.get('neu_min')
        rec.neu_max = request.POST.get('neu_max')

        rec.hgb = request.POST.get('hgb')
        rec.hgb_min = request.POST.get('hgb_min')
        rec.hgb_max = request.POST.get('hgb_max')

        rec.trom_plt = request.POST.get('trom_plt')
        rec.trom_plt_min = request.POST.get('trom_plt_min')
        rec.trom_plt = request.POST.get('trom_plt_max')

        rec.eos = request.POST.get('eos')
        rec.eos_min = request.POST.get('eos_min')
        rec.eos_max = request.POST.get('eos_max')

        rec.bas = request.POST.get('bas')
        rec.bas_min = request.POST.get('bas_min')
        rec.bas_max = request.POST.get('bas_max')

        rec.b_lymf = request.POST.get('b_lymf')
        rec.b_lymf_min = request.POST.get('b_lymf_min')
        rec.b_lymf_max = request.POST.get('b_lymf_max')

        rec.t_lymf = request.POST.get('t_lymf')
        rec.t_lymf_min = request.POST.get('t_lymf_min')
        rec.t_lymf_max = request.POST.get('t_lymf_max')

        rec.t_help = request.POST.get('t_help')
        rec.t_help_min = request.POST.get('t_help_min')
        rec.t_help_max = request.POST.get('t_help_max')

        rec.t_cyto_lymf = request.POST.get('t_cyto_lymf')
        rec.t_cyto_lymf_min = request.POST.get('t_cyto_lymf_min')
        rec.t_cyto_lymf_max = request.POST.get('t_cyto_lymf_max')

        rec.cd3_cd4_cd3_cd8 = request.POST.get('cd3_cd4_cd3_cd8')
        rec.cd3_cd4_cd3_cd8_min = request.POST.get('cd3_cd4_cd3_cd8_min')
        rec.cd3_cd4_cd3_cd8_max = request.POST.get('cd3_cd4_cd3_cd8_max')

        rec.nk_cells_cyto = request.POST.get('nk_cells_cyto')
        rec.nk_cells_cyto_min = request.POST.get('nk_cells_cyto_min')
        rec.nk_cells_cyto_max = request.POST.get('nk_cells_cyto_max')

        rec.common_nk_cells = request.POST.get('common_nk_cells')
        rec.common_nk_cells_min = request.POST.get('common_nk_cells_min')
        rec.common_nk_cells_max = request.POST.get('common_nk_cells_max')

        rec.circ_immu_comps = request.POST.get('circ_immu_comps')
        rec.circ_immu_comps_min = request.POST.get('circ_immu_comps_min')
        rec.circ_immu_comps_max = request.POST.get('circ_immu_comps_max')

        rec.hct_test_spon = request.POST.get('hct_test_spon')
        rec.hct_test_spon_min = request.POST.get('hct_test_spon_min')
        rec.hct_test_spon_max = request.POST.get('hct_test_spon_max')

        rec.hct_test_stim = request.POST.get('hct_test_stim')
        rec.hct_test_stim_min = request.POST.get('hct_test_stim_min')
        rec.hct_test_stim_max = request.POST.get('hct_test_stim_max')

        rec.cd3_ifny_stim = request.POST.get('cd3_ifny_stim')
        rec.cd3_ifny_stim_min = request.POST.get('cd3_ifny_stim_min')
        rec.cd3_ifny_stim_max = request.POST.get('cd3_ifny_stim_max')

        rec.cd3_ifny_spon = request.POST.get('cd3_ifny_spon')
        rec.cd3_ifny_spon_min = request.POST.get('cd3_ifny_spon_min')
        rec.cd3_ifny_spon_max = request.POST.get('cd3_ifny_spon_max')

        rec.cd3_tnfa_stim = request.POST.get('cd3_tnfa_stim')
        rec.cd3_tnfa_stim_min = request.POST.get('cd3_tnfa_stim_min')
        rec.cd3_tnfa_stim_max = request.POST.get('cd3_tnfa_stim_max')

        rec.cd3_tnfa_spon = request.POST.get('cd3_tnfa_spon')
        rec.cd3_tnfa_spon_min = request.POST.get('cd3_tnfa_spon_min')
        rec.cd3_tnfa_spon_max = request.POST.get('cd3_tnfa_spon_max')

        rec.cd3_il2_stim = request.POST.get('cd3_il2_stim')
        rec.cd3_il2_stim_min = request.POST.get('cd3_il2_stim_min')
        rec.cd3_il2_stim_max = request.POST.get('cd3_il2_stim_max')

        rec.cd3_il2_spon = request.POST.get('cd3_il2_spon')
        rec.cd3_il2_spon_min = request.POST.get('cd3_il2_spon_min')
        rec.cd3_il2_spon_max = request.POST.get('cd3_il2_spon_max')

        rec.recomendations = request.POST.get('recomendations')
        rec.save()
    return HttpResponseRedirect("/")

def CreateAllWordDocument(records):
    conn = sq.connect('db.sqlite3')
    cursor = conn.cursor()
    for i in records:
        ssq = f"""SELECT * from analysis_manualinput WHERE (name= '{i[3]}')"""
        cursor.execute(ssq)
        data = cursor.fetchall()
        conn.commit()
        print(i)
        surname = i[1]
        neu_lymf = []
        neu_cd3 = []
        neu_cd4 = []
        neu_cd8 = []
        lymf_cd19 = []
        cd19_cd4 = []
        cd19_cd8 = []
        interferon = []
        fno = []
        interlikin = []
        date = []
        y_clear = ()
        for y in data:
            arr = [y[30], y[31], y[32], y[33], y[34], y[35]]
            res = []
            for val in arr:
                if val is None:
                    res.append(1)
                else:
                    res.append(val)

            neu_lymf.append(round(float(y[15]) / float(y[13]), 2))
            neu_cd3.append(round(float(y[15]) / float(y[20]), 2))
            neu_cd4.append(round(float(y[15]) / float(y[22]), 2))
            neu_cd8.append(round(float(y[15]) / float(y[23]), 2))
            lymf_cd19.append(round(float(y[13]) / float(y[21]), 2))
            cd19_cd4.append(round(float(y[21]) / float(y[22]), 2))
            cd19_cd8.append(round(float(y[21]) / float(y[23]), 2))
            interferon.append(round((float(res[0]) / float(res[1])), 2))
            fno.append(round(float(res[3]) / float(res[2]), 2))
            interlikin.append(round(float(res[4]) / float(res[5]), 2))
            date.append(y[4])
            surname = y[2]
            y_clear = y

        print(len(data))
        if len(data) > 2:
            #создание таблицы если больше 2 анализов
            for i in range(len(data)):  # Цикл по индексам списка
                doc = Document()
                if i == 1:  # Если это первый элемент списка, пропустить его
                    doc = Document()

                    neu_lymf_1 = neu_lymf[0]
                    neu_lymf_2 = neu_lymf[1]

                    neu_cd3_1 = neu_cd3[0]
                    neu_cd3_2 = neu_cd3[1]

                    neu_cd4_1 = neu_cd4[0]
                    neu_cd4_2 = neu_cd4[1]

                    neu_cd8_1 = neu_cd8[0]
                    neu_cd8_2 = neu_cd8[1]

                    lymf_cd19_1 = lymf_cd19[0]
                    lymf_cd19_2 = lymf_cd19[1]

                    cd19_cd4_1 = cd19_cd4[0]
                    cd19_cd4_2 = cd19_cd4[1]

                    cd19_cd8_1 = cd19_cd8[0]
                    cd19_cd8_2 = cd19_cd8[1]

                    interferon_1 = interferon[0]
                    interferon_2 = interferon[1]

                    fno_1 = fno[0]
                    fno_2 = fno[1]

                    interlikin_1 = interlikin[0]
                    interlikin_2 = interlikin[1]

                    date_ = date[1]

                    date1 = date[0]
                    date2 = date[1]

                    data_ = (
                        ("NEU/LYMF", "1.67 - 1.80", neu_lymf_1, neu_lymf_2, CheckMinOrMax(neu_lymf_1, neu_lymf_2)),
                        ("NEU/CD3", "2.30 - 3.70", neu_cd3_1, neu_cd3_2, CheckMinOrMax(neu_cd3_1, neu_cd3_2)),
                        ("NEU/CD4", "9.50 - 12.30", neu_cd4_1, neu_cd4_2, CheckMinOrMax(neu_cd4_1, neu_cd4_2)),
                        ("NEU/CD8", "3.00  - 5.00", neu_cd8_1, neu_cd8_2, CheckMinOrMax(neu_cd8_1, neu_cd8_2)),
                        ("LYMF/CD19", "9.60 - 10.00", lymf_cd19_1, lymf_cd19_2, CheckMinOrMax(lymf_cd19_1, lymf_cd19_2)),
                        ("CD19/CD4", "0.50 - 0.80", cd19_cd4_1, cd19_cd4_2, CheckMinOrMax(cd19_cd4_1, cd19_cd4_2)),
                        ("CD19/CD8", "0.20 - 0.30", cd19_cd8_1, cd19_cd8_2, CheckMinOrMax(cd19_cd8_1, cd19_cd8_2)),
                        ("Интерферон", "80 - 120", interferon_1, interferon_2, CheckMinOrMax(interferon_1, interferon_2)),
                        ("ФНО", "80 - 120", fno_1, fno_2, CheckMinOrMax(fno_1, fno_2)),
                        ("Интерликин", "80 - 120", interlikin_1, interlikin_2, CheckMinOrMax(interlikin_1, interlikin_2))
                    )
                    CreateMainWordTabele(y, doc)
                    CreateWordTabeleTwoSesons(data_, date1, date2, doc)
                    CreateWordDocumentDeviations(y, doc)
                    doc.add_paragraph(f'Рекомендации/Заключения по Калькулятора рекомендаций : {CreateRecomendationWithCalculateRec(y_clear)}')
                    doc.save(
                        f'Печатные Формы/' + f'{surname} {date_}.docx')
                else:
                    if i == 0:
                        doc = Document()
                        CreateMainWordTabele(y, doc)
                        CreateWordDocumentDeviations(y, doc)
                        date_ = date[0]
                        doc.add_paragraph(f'Рекомендации/Заключения по Калькулятора рекомендаций : {CreateRecomendationWithCalculateRec(y_clear)}')
                        doc.save(
                            f'Печатные Формы/' + f'{surname} {date_}.docx')
                    else:

                        neu_lymf_slice_lst = neu_lymf[:i + 1]  # Создание среза списка от начала до i-го элемента
                        neu_lymf_ = round(sum(neu_lymf_slice_lst) / len(neu_lymf_slice_lst),2)  # Вычисление среднего значения в срезе

                        neu_cd3_slice_lst = neu_cd3[:i + 1]  # Создание среза списка от начала до i-го элемента
                        neu_cd3_ = round(sum(neu_cd3_slice_lst) / len(neu_cd3_slice_lst) ,2) # Вычисление среднего значения в срезе

                        neu_cd4_slice_lst = neu_cd4[:i + 1]  # Создание среза списка от начала до i-го элемента
                        neu_cd4_ = round(sum(neu_cd4_slice_lst) / len(neu_cd4_slice_lst),2)  # Вычисление среднего значения в срезе

                        neu_cd8_slice_lst = neu_cd8[:i + 1]  # Создание среза списка от начала до i-го элемента
                        neu_cd8_ = round(sum(neu_cd8_slice_lst) / len(neu_cd8_slice_lst),2)  # Вычисление среднего значения в срезе

                        lymf_cd19_slice_lst = lymf_cd19[:i + 1]  # Создание среза списка от начала до i-го элемента
                        lymf_cd19_ = round(sum(lymf_cd19_slice_lst) / len(lymf_cd19_slice_lst),2)  # Вычисление среднего значения в срезе

                        cd19_cd4_slice_lst = cd19_cd4[:i + 1]  # Создание среза списка от начала до i-го элемента
                        cd19_cd4_ = round(sum(cd19_cd4_slice_lst) / len(cd19_cd4_slice_lst),2)  # Вычисление среднего значения в срезе

                        cd19_cd8_slice_lst = cd19_cd8[:i + 1]  # Создание среза списка от начала до i-го элемента
                        cd19_cd8_ = round(sum(cd19_cd8_slice_lst) / len(cd19_cd8_slice_lst),2)  # Вычисление среднего значения в срезе

                        interferon_slice_lst = interferon[:i + 1]  # Создание среза списка от начала до i-го элемента
                        interferon_ = round(sum(interferon_slice_lst) / len(interferon_slice_lst),2)  # Вычисление среднего значения в срезе

                        fno_slice_lst = fno[:i + 1]  # Создание среза списка от начала до i-го элемента
                        fno_ = round(sum(fno_slice_lst) / len(fno_slice_lst),2)  # Вычисление среднего значения в срезе

                        interlikin_slice_lst = interlikin[:i + 1]  # Создание среза списка от начала до i-го элемента
                        interlikin_ = round(sum(interlikin_slice_lst) / len(interlikin_slice_lst),2)  # Вычисление среднего значения в срезе

                        date_ = date[i]

                        data_ = (
                            ("NEU/LYMF", "1.67 - 1.80", neu_lymf_, neu_lymf[i], CheckMinOrMax(neu_lymf_, neu_lymf[i])),
                            ("NEU/CD3", "2.30 - 3.70", neu_cd3_, neu_cd3[i], CheckMinOrMax(neu_cd3_, neu_cd3[i])),
                            ("NEU/CD4", "9.50 - 12.30", neu_cd4_, neu_cd4[i], CheckMinOrMax(neu_cd4_, neu_cd4[i])),
                            ("NEU/CD8", "3.00  - 5.00", neu_cd8_, neu_cd8[i], CheckMinOrMax(neu_cd8_, neu_cd8[i])),
                            ("LYMF/CD19", "9.60 - 10.00", lymf_cd19_, lymf_cd19[i], CheckMinOrMax(lymf_cd19_, lymf_cd19[i])),
                            ("CD19/CD4", "0.50 - 0.80", cd19_cd4_, cd19_cd4[i], CheckMinOrMax(cd19_cd4_, cd19_cd4[i])),
                            ("CD19/CD8", "0.20 - 0.30", cd19_cd8_, cd19_cd8[i], CheckMinOrMax(cd19_cd8_, cd19_cd8[i])),
                            ("Интерферон", "80 - 120", interferon_, interferon[i], CheckMinOrMax(interferon_, interferon[i])),
                            ("ФНО", "80 - 120", fno_, fno[-1], CheckMinOrMax(fno_, fno[i])),
                            ("Интерликин", "80 - 120", interlikin_, interlikin[i], CheckMinOrMax(interlikin_, interlikin[i]))
                        )
                        CreateMainWordTabele(y, doc)
                        doc.add_paragraph('')

                        CreateWordDocumentManySeasons(data_, date_, doc)
                        doc.add_paragraph('')
                        CreateWordDocumentDeviations(y, doc)
                        doc.add_paragraph(f'Рекомендации/Заключения по Калькулятора рекомендаций : {CreateRecomendationWithCalculateRec(y)}')
                        doc.save(f'Печатные Формы/' + f'{surname} {date_}.docx')
        if len(data) == 2:
            for i in range(len(data)):  # Цикл по индексам списка
                doc = Document()

                neu_lymf_1 = neu_lymf[0]
                neu_lymf_2 = neu_lymf[1]

                neu_cd3_1 = neu_cd3[0]
                neu_cd3_2 = neu_cd3[1]

                neu_cd4_1 = neu_cd4[0]
                neu_cd4_2 = neu_cd4[1]

                neu_cd8_1 = neu_cd8[0]
                neu_cd8_2 = neu_cd8[1]

                lymf_cd19_1 = lymf_cd19[0]
                lymf_cd19_2 = lymf_cd19[1]

                cd19_cd4_1 = cd19_cd4[0]
                cd19_cd4_2 = cd19_cd4[1]

                cd19_cd8_1 = cd19_cd8[0]
                cd19_cd8_2 = cd19_cd8[1]

                interferon_1 = interferon[0]
                interferon_2 = interferon[1]

                fno_1 = fno[0]
                fno_2 = fno[1]

                interlikin_1 = interlikin[0]
                interlikin_2 = interlikin[1]

                date_ = date[1]

                date1 = date[0]
                date2 = date[1]

                data_ = (
                    ("NEU/LYMF", "1.67 - 1.80", neu_lymf_1, neu_lymf_2, CheckMinOrMax(neu_lymf_1, neu_lymf_2)),
                    ("NEU/CD3", "2.30 - 3.70", neu_cd3_1, neu_cd3_2, CheckMinOrMax(neu_cd3_1, neu_cd3_2)),
                    ("NEU/CD4", "9.50 - 12.30", neu_cd4_1, neu_cd4_2, CheckMinOrMax(neu_cd4_1, neu_cd4_2)),
                    ("NEU/CD8", "3.00  - 5.00", neu_cd8_1, neu_cd8_2, CheckMinOrMax(neu_cd8_1, neu_cd8_2)),
                    ("LYMF/CD19", "9.60 - 10.00", lymf_cd19_1, lymf_cd19_2, CheckMinOrMax(lymf_cd19_1, lymf_cd19_2)),
                    ("CD19/CD4", "0.50 - 0.80", cd19_cd4_1, cd19_cd4_2, CheckMinOrMax(cd19_cd4_1, cd19_cd4_2)),
                    ("CD19/CD8", "0.20 - 0.30", cd19_cd8_1, cd19_cd8_2, CheckMinOrMax(cd19_cd8_1, cd19_cd8_2)),
                    ("Интерферон", "80 - 120", interferon_1, interferon_2, CheckMinOrMax(interferon_1, interferon_2)),
                    ("ФНО", "80 - 120", fno_1, fno_2, CheckMinOrMax(fno_1, fno_2)),
                    ("Интерликин", "80 - 120", interlikin_1, interlikin_2, CheckMinOrMax(interlikin_1, interlikin_2))
                )
                CreateMainWordTabele(y, doc)
                doc.add_paragraph('')
                CreateWordTabeleTwoSesons(data_, date1, date2, doc)
                doc.add_paragraph('')
                CreateWordDocumentDeviations(y, doc)
                doc.add_paragraph(f'Рекомендации/Заключения по Калькулятора рекомендаций : {CreateRecomendationWithCalculateRec(y_clear)}')
                doc.save(f'Печатные Формы/' + f'{surname} {date_}.docx')
        if len(data) < 2:
            doc = Document()
            CreateMainWordTabele(y_clear, doc)
            doc.add_paragraph('')
            CreateWordDocumentDeviations(y_clear, doc)
            date_ = date[0]
            doc.add_paragraph(f'Рекомендации/Заключения по Калькулятора рекомендаций : {CreateRecomendationWithCalculateRec(y_clear)}')
            doc.save(f'Печатные Формы/' + f'{surname}{date_}.docx')
def create_all_word_document_view(request):
    user = User.objects.filter(id=request.user.id)
    records = list((ManualInput.objects.all()).values_list())
    CreateAllWordDocument(records)

    if len(user) != 0:
        return redirect('docxForms')
    else:
        return redirect('login')
def CreateRecomendationWithCalculateRec(data_):

    name_age_Pacient = data_[5]
    name_Season_Pacient = data_[11]
    name_WBC_Pacient = data_[12]
    name_LYMF_Pacient = data_[13]
    name_MON_Pacient = data_[14]
    name_NEU_Pacient = data_[15]
    name_HGB_Pacient = data_[16]
    name_PLT_Pacient = data_[17]
    name_EOS_Pacient = data_[18]
    name_BAS_Pacient = data_[19]
    name_CD3_Pacient = data_[20]
    name_CD19_Pacient = data_[21]
    name_CD4_Pacient = data_[22]
    name_CD8_Pacient = data_[23]
    name_ratio_Pacient = data_[24]
    name_NK_Ob_Pacient = data_[25]
    name_NK_Cit_Pacient = data_[26]
    name_CIK_Pacient = data_[27]
    name_HCI_CO_Pacient = data_[28]
    name_HCI_CT_Pacient = data_[29]
    name_CD3IFN_CO_Pacient = data_[30]
    name_CD3IFN_CT_Pacient = data_[31]
    name_CD3INF_CO_Pacient = data_[32]
    name_CD3INF_CT_Pacient = data_[33]
    name_CD3IL2_CO_Pacient = data_[34]
    name_CD3IL2_CT_Pacient = data_[35]

    conn = sq.connect('db.sqlite3')
    cursor = conn.cursor()
    ssq = f"""SELECT * from analysis_recomendation"""
    cursor.execute(ssq)
    data = cursor.fetchall()
    conn.commit()

    for y in data:
        name_age =  CheckIsNoneAgeAndSeason(y[0],name_age_Pacient)
        name_Season = CheckIsNoneAgeAndSeason(y[1],name_Season_Pacient)
        name_WBC = CheckIsNone(y[2],y[3],y[4],name_WBC_Pacient)
        name_LYMF = CheckIsNone(y[5],y[6],y[7],name_LYMF_Pacient)
        name_MON = CheckIsNone(y[8],y[9],y[10],name_MON_Pacient)
        name_NEU = CheckIsNone(y[11],y[12],y[13],name_NEU_Pacient)
        name_HGB = CheckIsNone(y[14],y[15],y[16], name_HGB_Pacient)
        name_PLT = CheckIsNone(y[17],y[18],y[19],name_PLT_Pacient)
        name_EOS = CheckIsNone(y[20],y[21],y[22],name_EOS_Pacient)
        name_BAS = CheckIsNone(y[23],y[24],y[25],name_BAS_Pacient)
        name_CD3 = CheckIsNone(y[26],y[27],y[28],name_CD3_Pacient)
        name_CD19 = CheckIsNone(y[29],y[30],y[31],name_CD19_Pacient)
        name_CD4 = CheckIsNone(y[32],y[33],y[34],name_CD4_Pacient)
        name_CD8 = CheckIsNone(y[35],y[36],y[37],name_CD8_Pacient)
        name_ratio = CheckIsNone(y[38],y[39],y[40],name_ratio_Pacient)
        name_NK_Ob = CheckIsNone(y[41],y[42],y[43],name_NK_Ob_Pacient)
        name_NK_Cit = CheckIsNone(y[44],y[45],y[46],name_NK_Cit_Pacient)
        name_CIK = CheckIsNone(y[47],y[48],y[49],name_CIK_Pacient)
        name_HCI_CO = CheckIsNone(y[50],y[51],y[52],name_HCI_CO_Pacient)
        name_HCI_CT = CheckIsNone(y[53],y[54],y[55], name_HCI_CT_Pacient)
        name_CD3IFN_CO = CheckIsNone(y[56],y[57],y[58],name_CD3IFN_CO_Pacient)
        name_CD3IFN_CT = CheckIsNone(y[59],y[60],y[61],name_CD3IFN_CT_Pacient)
        name_CD3INF_CO = CheckIsNone(y[62],y[63],y[64],name_CD3INF_CO_Pacient)
        name_CD3INF_CT = CheckIsNone(y[65],y[66],y[67],name_CD3INF_CT_Pacient)
        name_CD3IL2_CO = CheckIsNone(y[68],y[69],y[70],name_CD3IL2_CO_Pacient)
        name_CD3IL2_CT = CheckIsNone(y[71],y[72],y[73], name_CD3IL2_CT_Pacient)

        if name_age and name_Season and name_WBC and name_LYMF and name_MON and name_NEU and name_HGB and  name_PLT and  name_EOS and  name_BAS and  name_CD3 and name_CD19 and name_CD4 and name_CD8 and name_ratio and  name_NK_Ob and   name_NK_Cit and   name_CIK and name_HCI_CO and name_HCI_CT and  name_CD3IFN_CO and   name_CD3IFN_CT and  name_CD3INF_CO and  name_CD3INF_CT and  name_CD3IL2_CO and  name_CD3IL2_CT:
            return y[74]
def CreateMainWordTabele(rows, doc):

    surname = rows[2]
    date = rows[4]
    arr = [rows[30], rows[31], rows[32], rows[33], rows[34], rows[35]]
    res = []
    for i in arr:
        if i is None:
            res.append(1)
        else:
            res.append(i)
    interferon = round(float(res[0]) / float(res[1]), 2)
    fno = round(float(res[3]) / float(res[2]), 2)
    interlikin = round(float(res[4]) / float(res[5]), 2)
    for i in arr:
        if i is None:
            res.append(' ')
        else:
            res.append(i)
    res_list = []
    for i in rows:
        if i == 'NULL' or i is None:
            res_list.append(' ')
        else:
            res_list.append(i)

    data = (
        ("Фамилия", str(res_list[2])),
        ("Пол", str(res_list[3])),
        ("Дата анализа", str(res_list[4])),
        ("Возраст", str(res_list[5])),
        ("Диагноз основной", str(res_list[6])),
        ("Диагноз сопутствующий", str(res_list[7])),
        ("1.Гены", str(res_list[8])),
        ("2.Гены", str(res_list[9])),
        ("3.Гены", str(res_list[10])),
        ("Сезон", str(res_list[11])),

        ("РЕЗУЛЬАТЫ ГЕМОТОЛОГИЧЕСКОГО ИССЛЕДОВАНИЯ", ""),
        ("Лейкоциты (WBC)", str(res_list[12])),
        ("Лимфоциты (LYMF)", str(res_list[13])),
        ("Моноциты (MON)", str(res_list[14])),
        ("Нейтрофилы (NEU)", str(res_list[15])),
        ("Эозинофилы (EOS)", str(res_list[18])),
        ("Базофилы (BAS)", str(res_list[19])),
        ("Гемоглобин (HGB)", str(res_list[16])),
        ("Тромбоциты (PLT)", str(res_list[17])),

        ("ИМУННЫЙ СТАТУС", ""),
        ("Общие T-лимфоциты (CD45+CD3+)", str(res_list[20])),
        ("Общие В-лимфоциты (CD45+CD19+)", str(res_list[21])),
        ("Т-хелперы (CD45+CD3+CD4+)", str(res_list[22])),
        ("Соотношение CD3+CD4+/CD3+CD8+  ", str(res_list[24])),
        ("Т-цитотоксические лимфоциты (CD45+CD3+СD8+)", str(res_list[23])),
        ("Циркулирующие иммунные комплексы", str(res_list[27])),
        ("Общие NK-клетки (CD45+CD3-CD16+56+) ", str(res_list[26])),
        ("NK-клетки цитолитические (CD45+CD3-CD16brightCD56dim) ", str(res_list[25])),
        ("HCI-тест(спонтанный)", str(res_list[28])),
        ("HCI-тест(стимулированый)", str(res_list[29])),

        ("ЦИТОКИНОВЫЙ СТАТУС", ""),
        ("CD3+IFNy+(стимулированный)", str(res_list[30])),
        ("CD3+IFNy+(спонтанный)", str(res_list[31])),
        ("Индекс (CD3+IFNy+(стимулированный)/CD3+IFNy+(спонтанный))", str(interferon)),

        ("CD3+TFNy+(стимулироанный)	 ", str(res_list[32])),
        ("CD3+TFNy+(спонтанный)	 ", str(res_list[33])),
        ("Индекс (CD3+TNFa+(стимулированный)/CD3+TNFa+(спонтанный))", str(fno)),

        ("CD3+IL2+(стимулированный)	 ", str(res_list[34])),
        ("CD3+IL2+(спонтанный)	 ", str(res_list[35])),
        ("Индекс (CD3+IL2+(стимулированный)/CD3+IL2+(спонтанный))", str(interlikin)),
    )
    neu_lymf =round(float(res_list[15])/ float(res_list[13]), 2)

    neu_cd3 = round(float(res_list[15]) / float(res_list[21]), 2)
    neu_cd4 = round(float(res_list[15]) / float(res_list[22]), 2)
    neu_cd8 = round(float(res_list[15]) / float(res_list[23]), 2)

    neu_cd19 = round(float(res_list[15]) / float(res_list[20]), 2)
    cd19_cd4 = round(float(res_list[20]) / float(res_list[22]), 2)
    cd19_cd8 = round(float(res_list[20]) / float(res_list[23]), 2)

    interlikin = round(float(res[4]) / float(res[5]), 2)
    interferon = round(float(res[0]) / float(res[1]), 2)
    fno = round(float(res[3]) / float(res[2]), 2)

    CreateGraphTcelllink(neu_cd3, neu_cd4, neu_cd8, neu_lymf, surname, date)
    CreateGraphBcelllink(neu_lymf, neu_cd19, cd19_cd4, cd19_cd8, surname, date)
    CreateGraphCytokinePairs(interlikin, interferon, fno, surname, date)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    for id, name in data:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = name

    doc.add_paragraph(' ')
    doc.add_paragraph(' ')
    doc.add_paragraph(f'График T-клеточного звена {surname}')
    picture1 = doc.add_picture(f'Графики/T-клеточный {surname} {date}.png', width=Inches(5))
    doc.add_paragraph(f'График B-клеточного звена {surname}')
    picture2 = doc.add_picture(f'Графики/B-клеточный {surname} {date}.png', width=Inches(5))
    doc.add_paragraph(f'График цитокиновых пар {surname}')
    picture3 = doc.add_picture(f'Графики/Цитокиновые пары {surname} {date}.png', width=Inches(5))

    if date == 'None':
        date = ''
def CreateWordDocumentManySeasons(data, date, doc):

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименование\nпоказателя'
    hdr_cells[1].text = 'Референтные\nзначения'
    hdr_cells[2].text = f'Среднее за\nпредыдущие значения'
    hdr_cells[3].text = f'Данные от {date}'
    hdr_cells[4].text = 'Итог сравнения'
    for row in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = str(row[1])
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4])
def CreateWordTabeleTwoSesons(data, date1, date2, doc):

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименование\nпоказателя'
    hdr_cells[1].text = 'Референтные\nзначения'
    hdr_cells[2].text = f'Данные от {date1}'
    hdr_cells[3].text = f'Данные от {date2}'
    hdr_cells[4].text = 'Итог сравнения'
    for row in data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[0])
        row_cells[1].text = str(row[1])
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4])
def CreateWordDocumentDeviations(rows, doc):
    date = rows[4]
    arr = [rows[30], rows[31], rows[32], rows[33], rows[34], rows[35]]
    res = []
    for i in arr:
        if i is None:
            res.append(1)
        else:
            res.append(i)
    interferon = round(float(res[0]) / float(res[1]), 2)
    fno = round(float(res[3]) / float(res[2]), 2)
    interlikin = round(float(res[4]) / float(res[5]), 2)

    neu_lymf = round (float(rows[15]) / float(rows[13]) ,2)
    neu_cd3 = round (float(rows[15]) / float(rows[20]) ,2)
    neu_cd4 = round (float(rows[15]) / float(rows[22]) ,2)
    neu_cd8 = round (float(rows[15]) / float(rows[23]) ,2)
    lymf_cd19 = round (float(rows[13]) / float(rows[21]) ,2)
    cd19_cd4 = round (float(rows[21]) / float(rows[22]) ,2)
    cd19_cd8 = round (float(rows[21])/ float(rows[23]) ,2)

    normal_values_max = [1.80, #NEU/LYMF
                         3.70, #NEU/CD3
                         12.30,#NEU/CD4
                         5.00, #NEU/CD8
                         10.00,#LYMF/CD19
                         0.80, #CD19/CD4
                         0.30,#CD19/CD8
                         120, #FNO
                         120, #INTERLIKIN
                         120] #INTERFERON

    normal_values_min = [1.67, #NEU/LYMF
                         2.30, #NEU/CD3
                         9.50, #NEU/CD4
                         3.00, #NEU/CD8
                         9.60, #LYMF/CD19
                         0.50, #CD19/CD4
                         0.20, #CD19/CD8
                         80,  # FNO
                         80,  # INTERLIKIN
                         80]  # INTERFERON


    doc.add_paragraph('Показатели T-клеточного иммунитета')
    doc.add_paragraph('     NEU/LYMF ' f'{CheckPercent(normal_values_max[0], normal_values_min[0], neu_lymf)} {CheckNumber(normal_values_max[0], normal_values_min[0], neu_lymf)}')
    doc.add_paragraph('     NEU/CD3 'f'{CheckPercent(normal_values_max[1], normal_values_min[1], neu_cd3)} {CheckNumber(normal_values_max[1], normal_values_min[1], neu_cd3)}')
    doc.add_paragraph('     NEU/CD4 'f'{CheckPercent(normal_values_max[2], normal_values_min[2], neu_cd4)} {CheckNumber(normal_values_max[2], normal_values_min[2], neu_cd4)}')
    doc.add_paragraph('     NEU/CD8 'f'{CheckPercent(normal_values_max[3], normal_values_min[3], neu_cd8)} {CheckNumber(normal_values_max[3], normal_values_min[3], neu_cd8)}')

    doc.add_paragraph('Показатели B-клеточного иммунитета')
    doc.add_paragraph('     LYMF/CD19 'f'{CheckPercent(normal_values_max[4], normal_values_min[4], lymf_cd19)} {CheckNumber(normal_values_max[4], normal_values_min[4], lymf_cd19)}')
    doc.add_paragraph('     CD19/CD4 'f'{CheckPercent(normal_values_max[5], normal_values_min[5], cd19_cd4)} {CheckNumber(normal_values_max[5], normal_values_min[5], cd19_cd4)}')
    doc.add_paragraph('     CD19/CD8 'f'{CheckPercent(normal_values_max[6], normal_values_min[6], cd19_cd8)} {CheckNumber(normal_values_max[6], normal_values_min[6], cd19_cd8)}')

    doc.add_paragraph('Цитокиновые пары')
    doc.add_paragraph('     ФНО 'f'{CheckPercent(normal_values_max[7], normal_values_min[7], fno)} {CheckNumber(normal_values_max[7], normal_values_min[7], fno)}')
    doc.add_paragraph('     ИНТЕРФЕРОН 'f'{CheckPercent(normal_values_max[8], normal_values_min[8], interferon)} {CheckNumber(normal_values_max[8], normal_values_min[8], interferon)}')
    doc.add_paragraph('     ИНТЕРЛЕКИН 'f'{CheckPercent(normal_values_max[9], normal_values_min[9], interlikin)} {CheckNumber(normal_values_max[9], normal_values_min[9], interlikin)}')
def CheckPercent(max, min, value):
    if value == 0:
        return 'N/A'
    elif max == 0 or min == 0:
        return 'N/A'
    else:
        maxProc = str(round(max / value, 2))
        minProc = str(round(min / value, 2))
        if (value > max):
            return f"- Отклонение от нормы больше {maxProc[:2]}% нормы вверх"
        if (value < min):
            return f"- Отклонение от нормы больше {minProc[:2]}% нормы вниз"
        if (value < min and value > max):
            return "- В пределах нормы"
def CheckNumber(max, min, value):
    if (value > max):
        return f'- Отклонение от нормы на {round(value-max,2)} нормы вверх'
    if (value < min):
        return f'- Отклонение от нормы на {round(value-min,2)} нормы вниз'
    if (value < min and value > max):
        return "- В пределах нормы"
def CheckMinOrMax(value, value2):
    if (value2 > value):
        return "Увеличился"
    if (value2 < value):
        return "Уменьшился"
    if (value2 < value and value2 > value):
        return "- В пределах нормы"
def CheckIsNone(value,value_min,value_max,value_pacient):
    #print(value,":", value_min,":", value_max,":", value_pacient)
    if value == "":
        if value_min == "":
            return True
        else:
            return value_min <= value_pacient <= value_max
    else:
        return value_pacient == value
def CheckIsNoneAgeAndSeason(value,value_pacient):
    if value == "":
       return True
    else:
        return value_pacient == value

